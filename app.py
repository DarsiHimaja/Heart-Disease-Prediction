from flask import Flask, request, jsonify, render_template, session
from flask_cors import CORS
import sqlite3
import joblib
import numpy as np
import os
import hashlib
import random
import re
from datetime import datetime, timedelta

app = Flask(__name__)
CORS(app, supports_credentials=True)
app.secret_key = 'your-secret-key-change-this-in-production'  # Change this!

# ---------- Database ----------
def init_db():
    conn = sqlite3.connect("heart.db")
    c = conn.cursor()
    
    # Users table
    c.execute('''CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT,
                    email TEXT UNIQUE,
                    password TEXT
                )''')
    
    # Predictions table
    c.execute('''CREATE TABLE IF NOT EXISTS predictions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    email TEXT,
                    risk_level TEXT,
                    prediction_result REAL
                )''')
    
    # Add created_at column if it doesn't exist
    try:
        c.execute("ALTER TABLE predictions ADD COLUMN created_at DATETIME")
    except sqlite3.OperationalError:
        pass  # Column already exists
    
    # Add health parameter columns if they don't exist
    health_columns = [
        "age INTEGER", "sex INTEGER", "cp INTEGER", "trestbps INTEGER", "chol INTEGER",
        "fbs INTEGER", "restecg INTEGER", "thalach INTEGER", "exang INTEGER",
        "oldpeak REAL", "slope INTEGER", "ca INTEGER", "thal INTEGER"
    ]
    
    for column in health_columns:
        try:
            c.execute(f"ALTER TABLE predictions ADD COLUMN {column}")
        except sqlite3.OperationalError:
            pass  # Column already exists
    
    conn.commit()
    conn.close()

init_db()

# ---------- Load Model ----------
model = joblib.load("heart_model.pkl")
scaler = joblib.load("scaler.pkl")

# ---------- Helper: Hash password ----------
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# ---------- Routes ----------
@app.route("/")
def home():
    return render_template("index.html")

@app.route('/favicon.ico')
def favicon():
    return '', 204  # No content response

# ---------- Register ----------
@app.route("/register", methods=["POST"])
def register():
    data = request.get_json()
    name = data.get("name")
    email = data.get("email")
    password = data.get("password")

    if not all([name, email, password]):
        return jsonify({"ok": False, "message": "All fields are required"}), 400

    conn = sqlite3.connect("heart.db")
    c = conn.cursor()
    try:
        c.execute("INSERT INTO users (name, email, password) VALUES (?, ?, ?)",
                  (name, email, hash_password(password)))
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({"ok": False, "message": "Email already registered"}), 400
    conn.close()
    return jsonify({"ok": True, "message": "Registered successfully"})

# ---------- Login ----------
@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    email = data.get("email")
    password = data.get("password")

    conn = sqlite3.connect("heart.db")
    c = conn.cursor()
    c.execute("SELECT name, password FROM users WHERE email=?", (email,))
    user = c.fetchone()
    conn.close()

    if user and user[1] == hash_password(password):
        # Store user info in session
        session['user_email'] = email
        session['user_name'] = user[0]
        session['logged_in'] = True
        return jsonify({"ok": True, "message": "Login successful", "user": {"email": email, "name": user[0]}})
    return jsonify({"ok": False, "message": "Invalid credentials"}), 401

# ---------- Logout ----------
@app.route("/logout", methods=["POST"])
def logout():
    session.clear()
    return jsonify({"ok": True, "message": "Logged out successfully"})

# ---------- Check Login Status ----------
@app.route("/check-auth", methods=["GET"])
def check_auth():
    if session.get('logged_in'):
        return jsonify({
            "ok": True, 
            "logged_in": True,
            "user": {
                "email": session.get('user_email'),
                "name": session.get('user_name')
            }
        })
    return jsonify({"ok": True, "logged_in": False})

# ---------- Predict ----------
@app.route("/predict", methods=["POST"])
def predict():
    data = request.get_json()
    email = data.get("user_email")

    values = [
        data.get("age", 0), data.get("sex", 0), data.get("cp", 0),
        data.get("trestbps", 0), data.get("chol", 0), data.get("fbs", 0),
        data.get("restecg", 0), data.get("thalach", 0), data.get("exang", 0),
        data.get("oldpeak", 0), data.get("slope", 0), data.get("ca", 0),
        data.get("thal", 0)
    ]

    print(f"Manual input values: {values}")  # Debug print
    
    arr = np.array([values])
    arr = scaler.transform(arr)
    pred = model.predict_proba(arr)[0][1] * 100
    
    print(f"Manual prediction: {pred}%")  # Debug print

    risk = "Low" if pred < 35 else "Moderate" if pred < 65 else "High"

    conn = sqlite3.connect("heart.db")
    c = conn.cursor()
    c.execute("""INSERT INTO predictions (email, risk_level, prediction_result, created_at, age, sex, cp, trestbps, chol, fbs, restecg, thalach, exang, oldpeak, slope, ca, thal) 
                 VALUES (?, ?, ?, datetime('now', 'localtime'), ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
              (email, risk, pred, data.get("age", 0), data.get("sex", 0), data.get("cp", 0), data.get("trestbps", 0), 
               data.get("chol", 0), data.get("fbs", 0), data.get("restecg", 0), data.get("thalach", 0), 
               data.get("exang", 0), data.get("oldpeak", 0), data.get("slope", 0), data.get("ca", 0), data.get("thal", 0)))
    conn.commit()
    conn.close()

    return jsonify({"ok": True, "result": {"risk_level": risk, "prediction_result": pred}})

# ---------- Get Predictions ----------
@app.route("/predictions")
def predictions():
    email = request.args.get("email")
    conn = sqlite3.connect("heart.db")
    c = conn.cursor()
    c.execute("SELECT * FROM predictions WHERE email=? ORDER BY id DESC", (email,))
    rows = []
    for r in c.fetchall():
        # If no timestamp, use a fixed time based on ID to avoid changing times
        timestamp = r[4] if len(r) > 4 and r[4] else f"2024-01-{(r[0] % 30) + 1:02d} 07:{(r[0] % 60):02d}:00"
        rows.append({
            "id": r[0], "email": r[1], "risk_level": r[2], "prediction_result": r[3], "created_at": timestamp,
            "age": r[5] if len(r) > 5 and r[5] is not None else 0,
            "sex": r[6] if len(r) > 6 and r[6] is not None else 0,
            "cp": r[7] if len(r) > 7 and r[7] is not None else 0,
            "trestbps": r[8] if len(r) > 8 and r[8] is not None else 0,
            "chol": r[9] if len(r) > 9 and r[9] is not None else 0,
            "fbs": r[10] if len(r) > 10 and r[10] is not None else 0,
            "restecg": r[11] if len(r) > 11 and r[11] is not None else 0,
            "thalach": r[12] if len(r) > 12 and r[12] is not None else 0,
            "exang": r[13] if len(r) > 13 and r[13] is not None else 0,
            "oldpeak": r[14] if len(r) > 14 and r[14] is not None else 0,
            "slope": r[15] if len(r) > 15 and r[15] is not None else 0,
            "ca": r[16] if len(r) > 16 and r[16] is not None else 0,
            "thal": r[17] if len(r) > 17 and r[17] is not None else 0
        })
    conn.close()
    return jsonify({"predictions": rows})

# ---------- Delete Prediction ----------
@app.route("/delete/<int:id>", methods=["DELETE"])
def delete(id):
    conn = sqlite3.connect("heart.db")
    c = conn.cursor()
    c.execute("DELETE FROM predictions WHERE id=?", (id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})

# ---------- Debug File Content ----------
@app.route("/debug-file", methods=["POST"])
def debug_file():
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({"error": "No file"})
        
        # Read file content directly
        content = file.read().decode('utf-8', errors='ignore')
        file.seek(0)  # Reset file pointer
        
        return jsonify({
            "filename": file.filename,
            "content": content,
            "length": len(content)
        })
    except Exception as e:
        return jsonify({"error": str(e)})

# ---------- OCR Implementation ----------
@app.route("/ocr", methods=["POST"])
def ocr():
    try:
        print("=== OCR REQUEST RECEIVED ===")
        
        file = request.files.get('file')
        email = request.form.get('user_email')
        
        if not file or not email:
            return jsonify({"ok": False, "message": "File and email required"}), 400
        
        print(f"Processing file: {file.filename}")
        file_ext = os.path.splitext(file.filename)[1].lower() if file.filename else '.txt'
        content = ""
        
        try:
            # Read file content based on type
            if file_ext == '.txt':
                content = file.read().decode('utf-8', errors='ignore')
                
            elif file_ext == '.docx':
                try:
                    from docx import Document
                    from io import BytesIO
                    
                    # Read DOCX from memory
                    file_stream = BytesIO(file.read())
                    doc = Document(file_stream)
                    
                    # Extract text from paragraphs
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    
                    # Extract text from tables
                    table_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    table_text.append(cell.text.strip())
                    
                    content = '\n'.join(paragraphs + table_text)
                    print(f"DOCX content extracted: {len(content)} characters")
                    
                except Exception as docx_error:
                    print(f"DOCX processing error: {docx_error}")
                    return jsonify({"ok": False, "message": f"Error processing DOCX file: {str(docx_error)}"}), 400
                    
            elif file_ext == '.pdf':
                try:
                    import PyPDF2
                    from io import BytesIO
                    
                    # Read PDF from memory
                    file_stream = BytesIO(file.read())
                    pdf_reader = PyPDF2.PdfReader(file_stream)
                    
                    # Extract text from all pages
                    pdf_text = []
                    for page in pdf_reader.pages:
                        pdf_text.append(page.extract_text())
                    
                    content = '\n'.join(pdf_text)
                    print(f"PDF content extracted: {len(content)} characters")
                    
                except Exception as pdf_error:
                    print(f"PDF processing error: {pdf_error}")
                    return jsonify({"ok": False, "message": f"Error processing PDF file: {str(pdf_error)}"}), 400
                    
            else:
                # Try as text file for other extensions
                content = file.read().decode('utf-8', errors='ignore')
                
        except Exception as read_error:
            print(f"File reading error: {read_error}")
            return jsonify({"ok": False, "message": f"Error reading file: {str(read_error)}"}), 400
        
        print(f"Content extracted: {len(content)} characters")
        print(f"First 200 chars: {repr(content[:200])}")
        
        if not content.strip():
            return jsonify({"ok": False, "message": "No content found in file. Please check your file."}), 400
        
        # Parse medical data
        extracted_data = parse_medical_text(content)
        
        if not extracted_data:
            return jsonify({
                "ok": False, 
                "message": "No medical data found. Please ensure your file contains medical parameters like Age, Gender, Blood Pressure, Cholesterol, etc.",
                "file_content_preview": content[:300]
            }), 400
        
        # Make prediction with extracted data
        defaults = {
            "age": 50, "sex": 1, "cp": 0, "trestbps": 120, "chol": 200,
            "fbs": 0, "restecg": 0, "thalach": 150, "exang": 0,
            "oldpeak": 1.0, "slope": 1, "ca": 0, "thal": 1
        }
        
        prediction_values = []
        for field in ["age", "sex", "cp", "trestbps", "chol", "fbs", "restecg", "thalach", "exang", "oldpeak", "slope", "ca", "thal"]:
            prediction_values.append(extracted_data.get(field, defaults[field]))
        
        # Make ML prediction
        arr = np.array([prediction_values])
        arr = scaler.transform(arr)
        pred = model.predict_proba(arr)[0][1] * 100
        risk = "Low" if pred < 35 else "Moderate" if pred < 65 else "High"
        
        # Save to database
        conn = sqlite3.connect("heart.db")
        c = conn.cursor()
        c.execute("""INSERT INTO predictions (email, risk_level, prediction_result, created_at, age, sex, cp, trestbps, chol, fbs, restecg, thalach, exang, oldpeak, slope, ca, thal) 
                     VALUES (?, ?, ?, datetime('now', 'localtime'), ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                  (email, risk, pred, *prediction_values))
        conn.commit()
        conn.close()
        
        return jsonify({
            "ok": True,
            "extracted": extracted_data,
            "file_content": content[:500],
            "result": {"risk_level": risk, "prediction_result": pred}
        })
        
    except Exception as e:
        print(f"OCR Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"ok": False, "message": f"Processing failed: {str(e)}"}), 500

# ---------- Parse Medical Text ----------
def parse_medical_text(text):
    """Extract ONLY values actually found in the file - no defaults"""
    try:
        print(f"=== RAW FILE CONTENT ===")
        print(text)
        print(f"=== END RAW CONTENT ===")
        
        extracted = {}
        found_count = 0
        
        # More flexible patterns for extraction
        
        # Extract Age - multiple patterns
        age_patterns = [
            r'Age\s*[:\-]?\s*(\d+)',
            r'age\s*[:\-]?\s*(\d+)',
            r'AGE\s*[:\-]?\s*(\d+)',
            r'(\d+)\s*years?\s*old',
            r'Patient.*?(\d+).*?years?'
        ]
        
        for pattern in age_patterns:
            age_match = re.search(pattern, text, re.IGNORECASE)
            if age_match:
                age_val = int(age_match.group(1))
                if 18 <= age_val <= 120:  # Reasonable age range
                    extracted["age"] = age_val
                    found_count += 1
                    print(f"✓ Age found: {extracted['age']}")
                    break
        
        # Extract Gender - multiple patterns
        gender_patterns = [
            r'Gender\s*[:\-]?\s*(\w+)',
            r'Sex\s*[:\-]?\s*(\w+)',
            r'gender\s*[:\-]?\s*(\w+)',
            r'sex\s*[:\-]?\s*(\w+)',
            r'(male|female|M|F)\b'
        ]
        
        for pattern in gender_patterns:
            gender_match = re.search(pattern, text, re.IGNORECASE)
            if gender_match:
                gender = gender_match.group(1).lower()
                if 'female' in gender or gender == 'f':
                    extracted["sex"] = 0
                elif 'male' in gender or gender == 'm':
                    extracted["sex"] = 1
                
                if "sex" in extracted:
                    found_count += 1
                    print(f"✓ Gender found: {gender} -> {extracted['sex']}")
                    break
        
        # Extract Blood Pressure - multiple patterns
        bp_patterns = [
            r'(?:Resting\s+)?Blood\s+Pressure\s*[:\-]?\s*(\d+)',
            r'BP\s*[:\-]?\s*(\d+)',
            r'Systolic\s*[:\-]?\s*(\d+)',
            r'(\d+)\s*mmHg',
            r'(\d+)/(\d+)\s*mmHg'  # Handle BP like 120/80
        ]
        
        for pattern in bp_patterns:
            bp_match = re.search(pattern, text, re.IGNORECASE)
            if bp_match:
                bp_val = int(bp_match.group(1))
                if 80 <= bp_val <= 250:  # Reasonable BP range
                    extracted["trestbps"] = bp_val
                    found_count += 1
                    print(f"✓ BP found: {extracted['trestbps']}")
                    break
        
        # Extract Cholesterol - multiple patterns
        chol_patterns = [
            r'Cholesterol\s*[:\-]?\s*(\d+)',
            r'cholesterol\s*[:\-]?\s*(\d+)',
            r'CHOL\s*[:\-]?\s*(\d+)',
            r'Total\s+Cholesterol\s*[:\-]?\s*(\d+)'
        ]
        
        for pattern in chol_patterns:
            chol_match = re.search(pattern, text, re.IGNORECASE)
            if chol_match:
                chol_val = int(chol_match.group(1))
                if 100 <= chol_val <= 600:  # Reasonable cholesterol range
                    extracted["chol"] = chol_val
                    found_count += 1
                    print(f"✓ Cholesterol found: {extracted['chol']}")
                    break
        
        # Extract Heart Rate - multiple patterns
        hr_patterns = [
            r'(?:Max\s+)?Heart\s+Rate\s*[:\-]?\s*(\d+)',
            r'Maximum\s+Heart\s+Rate\s*[:\-]?\s*(\d+)',
            r'HR\s*[:\-]?\s*(\d+)',
            r'(\d+)\s*bpm'
        ]
        
        for pattern in hr_patterns:
            hr_match = re.search(pattern, text, re.IGNORECASE)
            if hr_match:
                hr_val = int(hr_match.group(1))
                if 60 <= hr_val <= 220:  # Reasonable heart rate range
                    extracted["thalach"] = hr_val
                    found_count += 1
                    print(f"✓ Heart Rate found: {extracted['thalach']}")
                    break
        
        # Extract Chest Pain Type - multiple patterns
        cp_patterns = [
            r'Chest\s+Pain\s+Type\s*[:\-]?\s*([^\n\r]+)',
            r'chest\s+pain\s*[:\-]?\s*([^\n\r]+)',
            r'CP\s*[:\-]?\s*([^\n\r]+)'
        ]
        
        for pattern in cp_patterns:
            cp_match = re.search(pattern, text, re.IGNORECASE)
            if cp_match:
                cp_type = cp_match.group(1).strip().lower()
                if 'typical' in cp_type and 'angina' in cp_type:
                    extracted["cp"] = 0
                elif 'atypical' in cp_type:
                    extracted["cp"] = 1
                elif 'non-anginal' in cp_type or 'non anginal' in cp_type:
                    extracted["cp"] = 2
                elif 'asymptomatic' in cp_type:
                    extracted["cp"] = 3
                
                if "cp" in extracted:
                    found_count += 1
                    print(f"✓ Chest Pain found: {cp_type} -> {extracted['cp']}")
                    break
        
        # Extract Fasting Blood Sugar - multiple patterns
        fbs_patterns = [
            r'Fasting\s+Blood\s+Sugar\s*[:\-]?\s*(\w+)',
            r'FBS\s*[:\-]?\s*(\w+)',
            r'fasting\s+glucose\s*[:\-]?\s*(\w+)'
        ]
        
        for pattern in fbs_patterns:
            fbs_match = re.search(pattern, text, re.IGNORECASE)
            if fbs_match:
                fbs_val = fbs_match.group(1).lower()
                if 'yes' in fbs_val or 'positive' in fbs_val or 'high' in fbs_val:
                    extracted["fbs"] = 1
                elif 'no' in fbs_val or 'negative' in fbs_val or 'normal' in fbs_val:
                    extracted["fbs"] = 0
                
                if "fbs" in extracted:
                    found_count += 1
                    print(f"✓ FBS found: {fbs_val} -> {extracted['fbs']}")
                    break
        
        # Extract ECG - multiple patterns
        ecg_patterns = [
            r'(?:Rest\s+)?ECG\s*[:\-]?\s*([^\n\r]+)',
            r'(?:Resting\s+)?EKG\s*[:\-]?\s*([^\n\r]+)',
            r'electrocardiogram\s*[:\-]?\s*([^\n\r]+)'
        ]
        
        for pattern in ecg_patterns:
            ecg_match = re.search(pattern, text, re.IGNORECASE)
            if ecg_match:
                ecg_val = ecg_match.group(1).strip().lower()
                if 'normal' in ecg_val:
                    extracted["restecg"] = 0
                elif 'abnormal' in ecg_val or 'st-t' in ecg_val:
                    extracted["restecg"] = 1
                elif 'lvh' in ecg_val or 'hypertrophy' in ecg_val:
                    extracted["restecg"] = 2
                
                if "restecg" in extracted:
                    found_count += 1
                    print(f"✓ ECG found: {ecg_val} -> {extracted['restecg']}")
                    break
        
        # Extract Exercise Induced Angina - multiple patterns
        angina_patterns = [
            r'Exercise\s+Induced\s+Angina\s*[:\-]?\s*(\w+)',
            r'exercise\s+angina\s*[:\-]?\s*(\w+)',
            r'angina\s+on\s+exercise\s*[:\-]?\s*(\w+)'
        ]
        
        for pattern in angina_patterns:
            angina_match = re.search(pattern, text, re.IGNORECASE)
            if angina_match:
                angina_val = angina_match.group(1).lower()
                if 'yes' in angina_val or 'positive' in angina_val:
                    extracted["exang"] = 1
                elif 'no' in angina_val or 'negative' in angina_val:
                    extracted["exang"] = 0
                
                if "exang" in extracted:
                    found_count += 1
                    print(f"✓ Exercise Angina found: {angina_val} -> {extracted['exang']}")
                    break
        
        # Extract ST Depression - multiple patterns
        oldpeak_patterns = [
            r'ST\s+Depression\s*[:\-]?\s*([\d.]+)',
            r'oldpeak\s*[:\-]?\s*([\d.]+)',
            r'ST\s+segment\s+depression\s*[:\-]?\s*([\d.]+)'
        ]
        
        for pattern in oldpeak_patterns:
            oldpeak_match = re.search(pattern, text, re.IGNORECASE)
            if oldpeak_match:
                oldpeak_val = float(oldpeak_match.group(1))
                if 0 <= oldpeak_val <= 10:  # Reasonable range
                    extracted["oldpeak"] = oldpeak_val
                    found_count += 1
                    print(f"✓ ST Depression found: {extracted['oldpeak']}")
                    break
        
        # Extract Slope - multiple patterns
        slope_patterns = [
            r'Slope\s*[:\-]?\s*([^\n\r]+)',
            r'ST\s+slope\s*[:\-]?\s*([^\n\r]+)'
        ]
        
        for pattern in slope_patterns:
            slope_match = re.search(pattern, text, re.IGNORECASE)
            if slope_match:
                slope_val = slope_match.group(1).strip().lower()
                if 'upsloping' in slope_val or 'up' in slope_val:
                    extracted["slope"] = 0
                elif 'flat' in slope_val:
                    extracted["slope"] = 1
                elif 'downsloping' in slope_val or 'down' in slope_val:
                    extracted["slope"] = 2
                
                if "slope" in extracted:
                    found_count += 1
                    print(f"✓ Slope found: {slope_val} -> {extracted['slope']}")
                    break
        
        # Extract Number of Vessels - multiple patterns
        vessel_patterns = [
            r'(?:Number\s+of\s+)?Major\s+Vessels\s*[:\-]?\s*(\d+)',
            r'vessels\s*[:\-]?\s*(\d+)',
            r'CA\s*[:\-]?\s*(\d+)'
        ]
        
        for pattern in vessel_patterns:
            vessel_match = re.search(pattern, text, re.IGNORECASE)
            if vessel_match:
                vessel_val = int(vessel_match.group(1))
                if 0 <= vessel_val <= 4:  # Valid range
                    extracted["ca"] = vessel_val
                    found_count += 1
                    print(f"✓ Vessels found: {extracted['ca']}")
                    break
        
        # Extract Thalassemia - multiple patterns
        thal_patterns = [
            r'Thalassemia\s*[:\-]?\s*([^\n\r]+)',
            r'Thal\s*[:\-]?\s*([^\n\r]+)',
            r'thalassemia\s*[:\-]?\s*([^\n\r]+)'
        ]
        
        for pattern in thal_patterns:
            thal_match = re.search(pattern, text, re.IGNORECASE)
            if thal_match:
                thal_val = thal_match.group(1).strip().lower()
                print(f"Raw Thalassemia text found: '{thal_val}'")
                
                if 'normal' in thal_val:
                    extracted["thal"] = 1
                elif 'fixed' in thal_val:
                    extracted["thal"] = 2
                elif 'reversible' in thal_val:
                    extracted["thal"] = 3
                
                if "thal" in extracted:
                    found_count += 1
                    print(f"✓ Thalassemia found: {thal_val} -> {extracted['thal']}")
                    break
        
        print(f"=== EXTRACTION COMPLETE ===")
        print(f"Found {found_count} values: {extracted}")
        
        # Return if we found at least some values
        if found_count > 0:
            return extracted
        else:
            print("❌ No medical values found in file")
            return None
        
    except Exception as e:
        print(f"❌ Error parsing: {e}")
        return None

# Helper functions for text parsing
def extract_age(text):
    """Extract age from text"""
    match = re.search(r'Age:\s*(\d+)', text, re.IGNORECASE)
    if match:
        age = int(match.group(1))
        print(f"Found age: {age}")
        return age
    return 50

def extract_bp(text):
    """Extract blood pressure from text"""
    match = re.search(r'Resting Blood Pressure.*?:\s*(\d+)', text, re.IGNORECASE)
    if match:
        bp = int(match.group(1))
        print(f"Found BP: {bp}")
        return bp
    return 120

def extract_cholesterol(text):
    """Extract cholesterol from text"""
    match = re.search(r'Cholesterol.*?:\s*(\d+)', text, re.IGNORECASE)
    if match:
        chol = int(match.group(1))
        print(f"Found cholesterol: {chol}")
        return chol
    return 200

def extract_heart_rate(text):
    """Extract max heart rate from text"""
    match = re.search(r'Max Heart Rate:\s*(\d+)', text, re.IGNORECASE)
    if match:
        hr = int(match.group(1))
        print(f"Found heart rate: {hr}")
        return hr
    return 150

def extract_oldpeak(text):
    """Extract ST Depression from text"""
    match = re.search(r'ST Depression.*?:\s*(\d+(?:\.\d+)?)', text, re.IGNORECASE)
    if match:
        oldpeak = float(match.group(1))
        print(f"Found oldpeak: {oldpeak}")
        return oldpeak
    return 1.0

def extract_gender_from_text(text):
    """Extract gender from text"""
    match = re.search(r'Gender:\s*(\w+)', text, re.IGNORECASE)
    if match:
        gender = match.group(1).lower()
        print(f"Found gender: {gender}")
        if 'female' in gender:
            return 0
        else:
            return 1
    return 1

def extract_chest_pain_from_text(text):
    """Extract chest pain type from text"""
    match = re.search(r'Chest Pain Type:\s*([^\n]+)', text, re.IGNORECASE)
    if match:
        cp_type = match.group(1).lower()
        print(f"Found chest pain type: {cp_type}")
        if 'typical angina' in cp_type:
            return 0
        elif 'atypical' in cp_type:
            return 1
        elif 'non-anginal' in cp_type:
            return 2
        elif 'asymptomatic' in cp_type:
            return 3
    return 0

def extract_fbs_from_text(text):
    """Extract fasting blood sugar from text"""
    match = re.search(r'Fasting Blood Sugar.*?:\s*(\w+)', text, re.IGNORECASE)
    if match:
        fbs = match.group(1).lower()
        print(f"Found FBS: {fbs}")
        if 'yes' in fbs:
            return 1
        else:
            return 0
    return 1

def extract_ecg_from_text(text):
    """Extract ECG result from text"""
    match = re.search(r'Rest ECG:\s*([^\n]+)', text, re.IGNORECASE)
    if match:
        ecg = match.group(1).lower()
        print(f"Found ECG: {ecg}")
        if 'normal' in ecg:
            return 0
        elif 'st-t abnormality' in ecg or 'abnormality' in ecg:
            return 1
        elif 'lvh' in ecg:
            return 2
    return 1

def extract_angina_from_text(text):
    """Extract exercise induced angina from text"""
    match = re.search(r'Exercise Induced Angina:\s*(\w+)', text, re.IGNORECASE)
    if match:
        angina = match.group(1).lower()
        print(f"Found angina: {angina}")
        if 'yes' in angina:
            return 1
        else:
            return 0
    return 1

def extract_slope_from_text(text):
    """Extract slope of peak exercise ST segment"""
    match = re.search(r'Slope.*?:\s*([^\n]+)', text, re.IGNORECASE)
    if match:
        slope = match.group(1).lower()
        print(f"Found slope: {slope}")
        if 'upsloping' in slope:
            return 0
        elif 'flat' in slope:
            return 1
        elif 'downsloping' in slope:
            return 2
    return 2

def extract_vessels_from_text(text):
    """Extract number of major vessels colored by fluoroscopy"""
    match = re.search(r'Number of Major Vessels.*?:\s*(\d+)', text, re.IGNORECASE)
    if match:
        vessels = int(match.group(1))
        print(f"Found vessels: {vessels}")
        return vessels
    return 3

def extract_thal_from_text(text):
    """Extract thalassemia result"""
    match = re.search(r'Thalassemia:\s*([^\n]+)', text, re.IGNORECASE)
    if match:
        thal = match.group(1).lower()
        print(f"Found thalassemia: {thal}")
        if 'normal' in thal:
            return 1
        elif 'fixed defect' in thal:
            return 2
        elif 'reversible defect' in thal:
            return 3
    return 2

# ---------- Run ----------
if __name__ == "__main__":
    app.run(debug=True)

# For Vercel deployment - this is the WSGI application
vercel_app = app
